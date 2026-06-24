import re
from html import escape
from pathlib import Path

from docx import Document
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.lib.units import mm
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer

from app.core.config import settings
from app.models.artifact import GeneratedResume


def _export_directory(resume: GeneratedResume) -> Path:
    directory = Path(settings.local_upload_dir).resolve() / str(resume.user_id) / "exports"
    directory.mkdir(parents=True, exist_ok=True)
    return directory


def _clean_line(line: str) -> str:
    return re.sub(r"^[-#]+\s*", "", line).strip()


def export_resume_pdf(resume: GeneratedResume) -> str:
    destination = _export_directory(resume) / f"resume-{resume.id}.pdf"
    styles = getSampleStyleSheet()
    story = []
    for raw_line in resume.markdown_content.splitlines():
        line = _clean_line(raw_line)
        if not line:
            story.append(Spacer(1, 3 * mm))
            continue
        if raw_line.startswith("# "):
            style = styles["Heading1"]
        elif raw_line.startswith("## "):
            style = styles["Heading2"]
        else:
            style = styles["BodyText"]
        prefix = "• " if raw_line.startswith("- ") else ""
        story.append(Paragraph(f"{prefix}{escape(line)}", style))
        story.append(Spacer(1, 1.5 * mm))
    document = SimpleDocTemplate(
        str(destination), pagesize=A4, rightMargin=18 * mm, leftMargin=18 * mm,
        topMargin=16 * mm, bottomMargin=16 * mm,
    )
    document.build(story)
    return str(destination)


def export_resume_docx(resume: GeneratedResume) -> str:
    destination = _export_directory(resume) / f"resume-{resume.id}.docx"
    document = Document()
    for raw_line in resume.markdown_content.splitlines():
        line = _clean_line(raw_line)
        if raw_line.startswith("# "):
            document.add_heading(line, level=0)
        elif raw_line.startswith("## "):
            document.add_heading(line, level=1)
        elif raw_line.startswith("- "):
            document.add_paragraph(line, style="List Bullet")
        elif line:
            document.add_paragraph(line)
    document.save(destination)
    return str(destination)
